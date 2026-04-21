"""
Process uploaded files: CSV, PDF, and ZIP archives containing CSV/PDF files.
Extracts text content and returns structured data for AI analysis.
"""
import os
import csv
import io
import zipfile
import tempfile


def _has_meaningful_text(text: str) -> bool:
    """Check if extracted text contains actual words, not just separators."""
    stripped = text.replace("|", "").replace("-", "").replace(" ", "").replace("\n", "")
    return len(stripped) > 30


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file. Tries pdfplumber first, falls back to PyPDF2."""
    text = _extract_with_pdfplumber(file_path)
    if _has_meaningful_text(text):
        return text

    # Fallback to PyPDF2
    fallback = _extract_with_pypdf2(file_path)
    if _has_meaningful_text(fallback):
        return fallback

    # Return whichever got more content
    return fallback if len(fallback) > len(text) else text


def _extract_with_pdfplumber(file_path: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    import pdfplumber
    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            cells = [str(cell or "").strip() for cell in row]
                            if any(cells):
                                text_parts.append(" | ".join(cells))
    except Exception as e:
        print(f"Error extracting text with pdfplumber {file_path}: {e}")
    return "\n".join(text_parts)


def _extract_with_pypdf2(file_path: str) -> str:
    """Fallback PDF extraction using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting text with PyPDF2 {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word document (.docx) using python-docx."""
    import docx
    text_parts = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    text_parts.append(" | ".join(cells))
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
    return "\n".join(text_parts)


def parse_csv_to_items(file_path: str) -> list:
    """Parse a CSV file into a list of row dicts."""
    items = []
    try:
        with open(file_path, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                items.append(row)
    except Exception as e:
        print(f"Error reading CSV {file_path}: {e}")
    return items


def process_zip(zip_path: str, extract_dir: str) -> list:
    """
    Extract a ZIP file and return list of extracted file paths.
    Only extracts CSV and PDF files, ignores others.
    """
    extracted = []
    try:
        with zipfile.ZipFile(zip_path, "r") as zf:
            for name in zf.namelist():
                # Skip hidden files and directories
                if name.startswith("__MACOSX") or name.startswith("."):
                    continue
                lower = name.lower()
                if lower.endswith((".csv", ".pdf", ".doc", ".docx")):
                    zf.extract(name, extract_dir)
                    extracted.append(os.path.join(extract_dir, name))
    except Exception as e:
        print(f"Error extracting ZIP {zip_path}: {e}")
    return extracted


def process_uploaded_files(file_paths: list) -> dict:
    """
    Process a list of uploaded files (CSV, PDF, or extracted from ZIP).
    Returns: {
        "items": [...],          # structured CSV rows (if any)
        "pdf_text": "...",       # combined PDF text (if any)
        "file_summary": "..."    # summary of what was uploaded
    }
    """
    all_items = []
    all_pdf_text = []
    file_names = []
    warnings = []

    for path in file_paths:
        basename = os.path.basename(path)
        lower = path.lower()

        if lower.endswith(".csv"):
            items = parse_csv_to_items(path)
            all_items.extend(items)
            file_names.append(f"{basename} (CSV, {len(items)} rows)")

        elif lower.endswith(".pdf"):
            text = extract_text_from_pdf(path)
            stripped = text.strip()
            if len(stripped) > 50:
                all_pdf_text.append(f"--- Contents of {basename} ---\n{text}")
                file_names.append(f"{basename} (PDF, {len(stripped)} chars)")
            elif stripped:
                # Very little text — likely scanned
                all_pdf_text.append(f"--- Contents of {basename} ---\n{text}")
                file_names.append(f"{basename} (PDF, low text - possibly scanned)")
                warnings.append(f"{basename} appears to be a scanned document with very little extractable text. For best results, upload a text-based PDF or CSV.")
            else:
                file_names.append(f"{basename} (PDF, no extractable text - likely scanned)")
                warnings.append(f"{basename} appears to be a scanned/image-based PDF. No text could be extracted. Please upload a text-based PDF or CSV instead.")

    return {
        "items": all_items,
        "pdf_text": "\n\n".join(all_pdf_text) if all_pdf_text else "",
        "file_summary": "; ".join(file_names) if file_names else "No files processed",
        "warnings": warnings,
    }
